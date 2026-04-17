from docx import *



name = "output1.docx"

# 打开文档
doc = Document(r"D:\project\pyprj\pythonProject1\code\docx写入表格\output.docx")
print(len(doc.tables))

# table = doc.tables[3]

content = []

with open('input.txt','r',encoding='utf-8') as f:
    for line in f:
        # print(line)
        line = line.strip()
        if line:
            content.append(line)
print("列表长度：" + str(len(content)))
# content = '本次课程设计如同一场精心设计的学术远征，我踏入了医疗影像人工智能的崭新领域。项目的核心使命，是构建一个能够精准识别“新冠感染、病毒性肺炎、肺部阴影、正常肺部”四类影像的智能诊断模型。在为期数周的探索中，我以PyTorch为舟，以数据为桨，完成了六个阶段的完整航程。'


print(content)

# 第一步
# for i in range(60):
#     table.add_row()

# for i in range(len(content)):
#     contenti = content[i]
#     answer = []
#



doc.save(name)






# cell = table.cell(0, 0)
#
# # table.cell(0, 0).text = content
#
#
# paragraph = cell.paragraphs[0]
#
#
# if len(paragraph.runs) > 0:
#     paragraph.runs[0].text = content
#     for run in paragraph.runs[1:]:
#         run.text = ""
# else:
#     paragraph.text = content
#


