from docx import *


doc = Document(r"D:\project\pyprj\pythonProject1\code\docx表格实战2\output.docx")

print(doc.tables)

print(len(doc.tables))

sum = 0

table = doc.tables[len(doc.tables)-1]

#
# sum = int(input("请输入需要添加的行数："))
#
#
# for i in range(sum):
#     table.add_row()


for i in range(len(table.rows)):
    table.cell(i,0).text=''



doc.save("output.docx")